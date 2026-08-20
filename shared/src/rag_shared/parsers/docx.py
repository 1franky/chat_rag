from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

import docx

from rag_shared.models import TextBlock


def parse(path: Path) -> Iterator[TextBlock]:
    document = docx.Document(str(path))

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield TextBlock(text=text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            yield TextBlock(text="\n".join(rows), section="tabla")
